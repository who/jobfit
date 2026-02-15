"""Tests for DOCX resume parsing."""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from jobfit.errors import EXIT_INPUT
from jobfit.parse_docx import parse_docx


@pytest.fixture()
def tmp_dir(tmp_path: Path) -> Path:
    """Return a temporary directory."""
    return tmp_path


def create_docx(path: Path, paragraphs: list[str]) -> None:
    """Create a DOCX file with given paragraphs."""
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def create_docx_with_table(path: Path, rows: list[list[str]]) -> None:
    """Create a DOCX file with a table."""
    doc = docx.Document()
    doc.add_paragraph("Resume Header")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    doc.save(str(path))


class TestParseDocx:
    """DOCX parsing extracts text correctly."""

    def test_single_paragraph(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "resume.docx"
        create_docx(docx_path, ["Hello World, this is a test resume."])
        result = parse_docx(docx_path)
        assert "Hello World" in result
        assert "test resume" in result

    def test_multiple_paragraphs(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "multi.docx"
        paras = ["First paragraph", "Second paragraph", "Third paragraph"]
        create_docx(docx_path, paras)
        result = parse_docx(docx_path)
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result

    def test_paragraphs_separated_by_newlines(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "newlines.docx"
        create_docx(docx_path, ["Para one", "Para two"])
        result = parse_docx(docx_path)
        assert "Para one\nPara two" in result

    def test_table_extraction(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "table.docx"
        rows = [["Skill", "Level"], ["Python", "Expert"], ["SQL", "Advanced"]]
        create_docx_with_table(docx_path, rows)
        result = parse_docx(docx_path)
        assert "Python" in result
        assert "Expert" in result
        assert "SQL" in result

    def test_file_not_found(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "nonexistent.docx"
        with pytest.raises(SystemExit) as exc_info:
            parse_docx(docx_path)
        assert exc_info.value.code == EXIT_INPUT

    def test_invalid_file(self, tmp_dir: Path) -> None:
        bad_path = tmp_dir / "notadocx.docx"
        bad_path.write_text("this is not a docx")
        with pytest.raises(SystemExit) as exc_info:
            parse_docx(bad_path)
        assert exc_info.value.code == EXIT_INPUT

    def test_empty_docx(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "empty.docx"
        doc = docx.Document()
        doc.save(str(docx_path))
        with pytest.raises(SystemExit) as exc_info:
            parse_docx(docx_path)
        assert exc_info.value.code == EXIT_INPUT

    def test_returns_string(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "type.docx"
        create_docx(docx_path, ["Some text here"])
        result = parse_docx(docx_path)
        assert isinstance(result, str)

    def test_word_count_from_result(self, tmp_dir: Path) -> None:
        docx_path = tmp_dir / "words.docx"
        create_docx(docx_path, ["one two three four five"])
        result = parse_docx(docx_path)
        words = result.split()
        assert len(words) >= 5
