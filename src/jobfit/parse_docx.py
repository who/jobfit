"""DOCX resume parsing using python-docx."""

from __future__ import annotations

from pathlib import Path

import docx

from jobfit.errors import EXIT_INPUT, print_error


def parse_docx(path: Path) -> str:
    """Extract text content from a DOCX file.

    Args:
        path: Path to the DOCX file.

    Returns:
        Extracted text as a string.

    Raises:
        SystemExit: If the file cannot be read or contains no text.
    """
    if not path.exists():
        print_error(
            f"Resume file not found: {path}",
            hint="Check the file path and try again.",
            exit_code=EXIT_INPUT,
        )

    try:
        doc = docx.Document(str(path))
    except Exception:
        print_error(
            f"Cannot read DOCX file: {path}",
            detail="The file is corrupted or not a valid DOCX document.",
            hint="Ensure the file is a valid, non-corrupted DOCX document.",
            exit_code=EXIT_INPUT,
        )

    paragraphs: list[str] = []

    # Extract paragraph text
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract table text (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append("\t".join(cells))

    full_text = "\n".join(paragraphs)

    if not full_text.strip():
        print_error(
            f"No text content found in DOCX: {path}",
            detail="The DOCX file appears to contain no extractable text.",
            hint="Ensure the document contains text content.",
            exit_code=EXIT_INPUT,
        )

    return full_text
